from tkinter import Tk, Label, Button, Entry, Text, messagebox, filedialog
from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.shared import Inches
from docx.shared import Cm
from docx.oxml.ns import nsdecls
from docx.oxml import parse_xml
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_ALIGN_VERTICAL
from docx2pdf import convert

image_path = None  # Initialize global variable for image path

def validate_name(name):
    return all(char.isalpha() or char.isspace() for char in name)

def validate_contact(contact):
    return contact.isdigit()

def save_resume():
    name = name_entry.get()
    contact = contact_entry.get()
    email = email_entry.get()
    education = education_entry.get("1.0", "end-1c")
    experience = experience_entry.get("1.0", "end-1c")
    skills = skills_entry.get("1.0", "end-1c")
    achievements = achievements_entry.get("1.0", "end-1c")
    about = about_entry.get("1.0", "end-1c")

    # Input validation
    if not (name and contact and email and education and experience and skills and achievements and about):
        messagebox.showerror("Error", "Please fill in all the required fields.")
        return
    elif not validate_name(name):
        messagebox.showerror("Error", "Name should contain only alphabets and spaces.")
        return
    elif not validate_contact(contact):
        messagebox.showerror("Error", "Contact should contain only numbers.")
        return

    try:
        # Get save path
        save_path = filedialog.asksaveasfilename(defaultextension=".pdf", filetypes=[("PDF Files", "*.pdf")])

        if save_path:
            # Get image path
            global image_path
            if image_path is None:
                messagebox.showerror("Error", "Please select an image.")
                return

            # Create DOCX
            docx_file = save_path.replace('.pdf', '.docx')
            create_docx_resume(docx_file, name, contact, email, education, experience, skills, achievements, about, image_path)

            # Convert DOCX to PDF
            convert(docx_file)

            messagebox.showinfo("Success", "Resume saved successfully as PDF.")
    except Exception as e:
        messagebox.showerror("Error", f"An error occurred: {str(e)}")

def create_docx_resume(file_name, name, contact, email, education, experience, skills, achievements, about, image_path):
    document = Document()

    # Adding header
    header = document.add_heading(level=1)
    header_run = header.add_run("Resume")
    header_run.bold = True
    header.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Adding separation line
    document.add_paragraph("----"*25)

    # Adding table for layout
    table = document.add_table(rows=1, cols=2)
    table.autofit = True
    table.style = document.styles['Table Grid']

    # Add image to table
    if image_path:
        cell = table.cell(0, 0)
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        cell.width = Inches(2)  # Adjust the width of the cell as needed
        paragraph = cell.paragraphs[0]
        run = paragraph.add_run()
        run.add_picture(image_path, width=Inches(1.25))  # Adjust the width as needed

    # Add personal information to table
    cell = table.cell(0, 1)
    cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
    personal_info = cell.paragraphs[0]
    personal_info.add_run(f"Name: {name}\n").bold = True
    personal_info.add_run(f"Email: {email}\n").bold = True
    personal_info.add_run(f"Phone: {contact}\n").bold = True

    # Adding education section
    education_heading = document.add_heading("Education", level=2)
    education_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_list_to_docx(document, education, center=True)

    # Adding experience section
    experience_heading = document.add_heading("Experience", level=2)
    experience_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_list_to_docx(document, experience, center=True)

    # Adding skills section
    skills_heading = document.add_heading("Skills", level=2)
    skills_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_list_to_docx(document, skills, center=True)

    # Adding achievements section
    achievements_heading = document.add_heading("Achievements", level=2)
    achievements_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_list_to_docx(document, achievements, center=True)

    # Adding about secti
    about_heading = document.add_heading("About", level=2)
    about_heading.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    add_list_to_docx(document, about, center=True)


    # Adjusting font size
    for paragraph in document.paragraphs:
        for run in paragraph.runs:
            run.font.size = Pt(12)

    # Save the document
    document.save(file_name)

def add_list_to_docx(document, content, center=False):
    paragraphs = content.split('\n')
    for item in paragraphs:
        if center:
            p = document.add_paragraph(item, style='List Bullet')
            p.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
        else:
            p = document.add_paragraph(item, style='List Bullet')

##        document.add_paragraph(item, style=document.styles['List Bullet'])

def select_image():
    global image_path
    image_path = filedialog.askopenfilename(title="Select Image")
    if image_path:
        messagebox.showinfo("Success", "Image selected successfully.")

# Create main window
root = Tk()
root.title("Resume Builder")
root.configure(bg='lightpink')  # Set background color

# Labels
Label(root, text="Name:", bg='lightgray', fg='black').grid(row=0, column=0, padx=5, pady=5)
Label(root, text="Contact no.:", bg='lightgray', fg='black').grid(row=1, column=0, padx=5, pady=5)
Label(root, text="Email:", bg='lightgray', fg='black').grid(row=2, column=0, padx=5, pady=5)
Label(root, text="Education:", bg='lightgray', fg='black').grid(row=3, column=0, padx=5, pady=5)
Label(root, text="Experience:", bg='lightgray', fg='black').grid(row=3, column=4, padx=5, pady=5)
Label(root, text="Skills:", bg='lightgray', fg='black').grid(row=4, column=0, padx=5, pady=5)
Label(root, text="Achievements:", bg='lightgray', fg='black').grid(row=4, column=4, padx=5, pady=5)
Label(root, text="About:", bg='lightgray', fg='black').grid(row=5, column=0, padx=5, pady=5)

# Entries
name_entry = Entry(root, bg='beige', validate='key')
name_entry.grid(row=0, column=1, padx=5, pady=5)
name_entry['validatecommand'] = (name_entry.register(validate_name), '%P')

contact_entry = Entry(root, bg='beige', validate='key')
contact_entry.grid(row=1, column=1, padx=5, pady=5)
contact_entry['validatecommand'] = (contact_entry.register(validate_contact), '%P')

email_entry = Entry(root, bg='beige')
email_entry.grid(row=2, column=1, padx=5, pady=5)

education_entry = Text(root, height=5, width=30, bg='beige')
education_entry.grid(row=3, column=1, padx=5, pady=5)

experience_entry = Text(root, height=5, width=30, bg='beige')
experience_entry.grid(row=3, column=5, padx=5, pady=5)

skills_entry = Text(root, height=5, width=30, bg='beige')
skills_entry.grid(row=4, column=1, padx=5, pady=5)

achievements_entry = Text(root, height=5, width=30, bg='beige')
achievements_entry.grid(row=4, column=5, padx=5, pady=5)

about_entry = Text(root, height=5, width=30, bg='beige')
about_entry.grid(row=5, column=1, padx=5, pady=5)

# Button to create resume
Button(root, text="Create Resume", command=save_resume, bg='lightgray', fg='black').grid(row=8, column=0, columnspan=2, padx=5, pady=10)
Button(root, text="Select Image", command=select_image, bg='lightgray', fg='black').grid(row=5, column=5, padx=5, pady=10)

root.mainloop()

